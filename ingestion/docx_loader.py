"""
DOCX Loader — extracts text with heading-aware section tracking.

Uses python-docx directly (rather than Docx2txtLoader) to iterate paragraphs
and track the nearest preceding heading, providing section_title metadata
for citations.
"""

import io
from typing import Union

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import CHUNK_SIZE, CHUNK_OVERLAP


def load_docx(
    source: Union[str, io.BytesIO],
    filename: str = "unknown.docx",
) -> list[Document]:
    """
    Load a DOCX file and return chunked Documents with section/paragraph metadata.

    Parameters
    ----------
    source : str | BytesIO
        File path or in-memory buffer.
    filename : str
        Original filename for metadata.

    Returns
    -------
    list[Document]
        Each chunk carries metadata:
        {"source_filename", "file_type": "docx", "section_title", "paragraph_index"}.
    """
    doc = DocxDocument(source)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    # ── Group paragraphs by the nearest preceding heading ──────────────
    sections: list[dict] = []  # [{"title": str, "text": str, "start_para": int}]
    current_title = "Untitled Section"
    current_text_parts: list[str] = []
    current_start_para = 0

    for i, para in enumerate(doc.paragraphs):
        style_name = (para.style.name or "").lower()

        # Detect headings (Heading 1, Heading 2, etc.)
        if style_name.startswith("heading") and para.text.strip():
            # Save the previous section
            if current_text_parts:
                sections.append(
                    {
                        "title": current_title,
                        "text": "\n".join(current_text_parts),
                        "start_para": current_start_para,
                    }
                )
            current_title = para.text.strip()
            current_text_parts = []
            current_start_para = i
        else:
            text = para.text.strip()
            if text:
                current_text_parts.append(text)

    # Don't forget the last section
    if current_text_parts:
        sections.append(
            {
                "title": current_title,
                "text": "\n".join(current_text_parts),
                "start_para": current_start_para,
            }
        )

    # ── Chunk each section ─────────────────────────────────────────────
    documents: list[Document] = []
    for section in sections:
        chunks = splitter.split_text(section["text"])
        for chunk_idx, chunk in enumerate(chunks):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source_filename": filename,
                        "file_type": "docx",
                        "section_title": section["title"],
                        "paragraph_index": section["start_para"] + chunk_idx,
                    },
                )
            )

    return documents
