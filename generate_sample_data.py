"""
Generate sample data files for testing FilingIQ.

Creates synthetic financial documents in all 5 supported formats:
  - PDF (10-K style MD&A and Risk Factors)
  - DOCX (Analyst memo with headings)
  - XLSX (Income statement, balance sheet)
  - CSV (Historical financial ratios)
  - TXT (Earnings call notes)

Run this script once to populate the data/ directory:
    python generate_sample_data.py
"""

import os
import sys

DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(DATA_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════
# 1. TXT — Earnings call notes (no dependencies needed)
# ═══════════════════════════════════════════════════════════════════════════

def create_sample_txt():
    content = """APPLE INC. Q4 2024 EARNINGS CALL NOTES
========================================

Date: October 31, 2024
Participants: Tim Cook (CEO), Luca Maestri (CFO), Analysts

KEY HIGHLIGHTS:
- Total revenue for Q4 2024 was $94.9 billion, up 6% year over year.
- iPhone revenue was $46.2 billion, representing a 6% increase from Q4 2023.
- Services revenue reached a new all-time high of $25.0 billion, up 12% year over year.
- Mac revenue was $7.7 billion, down 2% compared to Q4 2023.
- iPad revenue was $6.9 billion, a decline of 8% year over year.
- Wearables, Home and Accessories revenue was $9.0 billion, up 3%.

MANAGEMENT COMMENTARY:
Tim Cook emphasized the strong performance of the iPhone 16 lineup, noting that
customer satisfaction remains above 98%. The services ecosystem continues to expand
with over 1 billion paid subscriptions across Apple's platforms.

Luca Maestri highlighted that operating cash flow for the quarter was $26.8 billion.
The company returned over $29 billion to shareholders through dividends and share
repurchases during Q4. Gross margin for the quarter was 46.2%, up from 45.2% in
the year-ago quarter.

RISKS DISCUSSED:
- Foreign exchange headwinds expected to continue impacting international revenue.
- Supply chain diversification ongoing, with increased manufacturing in India and Vietnam.
- Regulatory challenges in the European Union related to the Digital Markets Act.
- Macroeconomic uncertainty in China impacting consumer spending patterns.

GUIDANCE:
For Q1 2025 (fiscal), Apple expects revenue between $123 billion and $127 billion.
Gross margin is expected to be between 46.0% and 47.0%.
Operating expenses are expected to be between $15.3 billion and $15.5 billion.

ANALYST Q&A HIGHLIGHTS:
- Morgan Stanley asked about AI integration: Tim Cook confirmed significant
  investments in Apple Intelligence features rolling out through 2025.
- Goldman Sachs inquired about capital allocation: Luca Maestri reaffirmed
  the goal of being net cash neutral over time.
- JP Morgan asked about India manufacturing: Tim Cook noted India now
  represents approximately 14% of iPhone production, up from 7% last year.
"""
    path = os.path.join(DATA_DIR, "apple_q4_2024_earnings_notes.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  ✅ Created: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# 2. CSV — Historical financial ratios
# ═══════════════════════════════════════════════════════════════════════════

def create_sample_csv():
    import csv

    path = os.path.join(DATA_DIR, "tech_company_ratios.csv")
    headers = [
        "Company", "Year", "Revenue_Billions", "Net_Income_Billions",
        "Gross_Margin_Pct", "Operating_Margin_Pct", "ROE_Pct",
        "Debt_to_Equity", "PE_Ratio", "EPS_Dollars"
    ]
    rows = [
        ["Apple", 2024, 394.3, 101.2, 46.2, 31.5, 157.4, 1.87, 30.2, 6.57],
        ["Apple", 2023, 383.3, 97.0, 44.1, 29.8, 171.9, 1.79, 29.1, 6.13],
        ["Apple", 2022, 394.3, 99.8, 43.3, 30.3, 196.9, 2.39, 24.8, 6.11],
        ["Apple", 2021, 365.8, 94.7, 41.8, 29.8, 150.1, 1.99, 28.7, 5.67],
        ["Microsoft", 2024, 245.1, 88.1, 69.4, 44.6, 37.4, 0.42, 35.1, 11.86],
        ["Microsoft", 2023, 211.9, 72.4, 68.9, 41.2, 35.1, 0.47, 32.4, 9.68],
        ["Microsoft", 2022, 198.3, 72.7, 68.4, 42.1, 43.7, 0.50, 28.1, 9.65],
        ["Microsoft", 2021, 168.1, 61.3, 68.9, 41.6, 47.1, 0.53, 35.6, 8.05],
        ["Tesla", 2024, 97.7, 7.1, 18.2, 7.2, 8.9, 0.11, 62.4, 2.20],
        ["Tesla", 2023, 96.8, 15.0, 18.2, 8.9, 22.5, 0.08, 73.1, 4.31],
        ["Tesla", 2022, 81.5, 12.6, 25.6, 16.8, 28.1, 0.07, 42.3, 3.62],
        ["Tesla", 2021, 53.8, 5.5, 25.3, 12.1, 15.1, 0.23, 112.8, 1.59],
    ]

    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)
    print(f"  ✅ Created: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# 3. XLSX — Income statement and balance sheet
# ═══════════════════════════════════════════════════════════════════════════

def create_sample_xlsx():
    import pandas as pd

    path = os.path.join(DATA_DIR, "apple_financials_2024.xlsx")

    income_data = {
        "Line Item": [
            "Total Revenue", "Cost of Revenue", "Gross Profit",
            "Research and Development", "Selling, General & Admin",
            "Total Operating Expenses", "Operating Income",
            "Interest Income", "Interest Expense",
            "Other Income/Expense", "Income Before Tax",
            "Income Tax Expense", "Net Income",
        ],
        "FY2024 ($M)": [
            394328, 212035, 182293,
            31370, 26146, 269551,
            124777, 5820, 3478,
            -415, 126704, 25488, 101216,
        ],
        "FY2023 ($M)": [
            383285, 214137, 169148,
            29915, 24932, 268984,
            114301, 5346, 3933,
            -565, 115149, 18129, 97020,
        ],
        "YoY Change (%)": [
            2.9, -1.0, 7.8,
            4.9, 4.9, 0.2,
            9.2, 8.9, -11.6,
            -26.5, 10.0, 40.6, 4.3,
        ],
    }

    balance_data = {
        "Line Item": [
            "Cash and Cash Equivalents", "Short-term Investments",
            "Accounts Receivable", "Inventories",
            "Total Current Assets", "Long-term Investments",
            "Property, Plant & Equipment", "Goodwill",
            "Total Assets",
            "Accounts Payable", "Short-term Debt",
            "Total Current Liabilities", "Long-term Debt",
            "Total Liabilities",
            "Common Stock", "Retained Earnings",
            "Total Stockholders' Equity",
        ],
        "FY2024 ($M)": [
            29943, 35228, 66243, 7286,
            152987, 100544, 44856, 0,
            364980,
            68960, 20855, 153981, 96807,
            308030,
            83276, -26233, 56950,
        ],
        "FY2023 ($M)": [
            29965, 31590, 60932, 6331,
            143566, 100544, 43715, 0,
            352583,
            62611, 18196, 145308, 95281,
            290437,
            73812, -11598, 62146,
        ],
    }

    cash_flow_data = {
        "Line Item": [
            "Net Income", "Depreciation & Amortization",
            "Stock-Based Compensation",
            "Changes in Working Capital",
            "Operating Cash Flow",
            "Capital Expenditures", "Acquisitions",
            "Investing Cash Flow",
            "Debt Repayment", "Share Buybacks",
            "Dividends Paid",
            "Financing Cash Flow",
        ],
        "FY2024 ($M)": [
            101216, 11519, 11688, 3430, 118253,
            -9959, -2100, -8342,
            -11400, -94949, -15234, -121583,
        ],
        "FY2023 ($M)": [
            97020, 11519, 10833, 1996, 110543,
            -10959, -1784, -7077,
            -9900, -77550, -15025, -108488,
        ],
    }

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame(income_data).to_excel(
            writer, sheet_name="Income Statement", index=False
        )
        pd.DataFrame(balance_data).to_excel(
            writer, sheet_name="Balance Sheet", index=False
        )
        pd.DataFrame(cash_flow_data).to_excel(
            writer, sheet_name="Cash Flow", index=False
        )

    print(f"  ✅ Created: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# 4. DOCX — Analyst memo
# ═══════════════════════════════════════════════════════════════════════════

def create_sample_docx():
    from docx import Document as DocxDoc
    from docx.shared import Inches, Pt

    path = os.path.join(DATA_DIR, "apple_analyst_memo_2024.docx")
    doc = DocxDoc()

    # Title
    doc.add_heading("Apple Inc. — Investment Research Memo", level=0)
    doc.add_paragraph("Date: November 15, 2024 | Analyst: Jane Smith, CFA")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Apple Inc. (AAPL) delivered solid fiscal year 2024 results, with total "
        "revenue of $394.3 billion, up 2.9% year over year. The company's services "
        "segment continues to be the primary growth driver, reaching $100 billion in "
        "annual revenue for the first time. We maintain our OVERWEIGHT rating with a "
        "price target of $245, representing 15% upside from current levels."
    )

    # Revenue Analysis
    doc.add_heading("Revenue Analysis", level=1)
    doc.add_paragraph(
        "iPhone revenue of $201.2 billion comprised 51% of total revenue, roughly "
        "flat year over year. The iPhone 16 launch in September showed strong initial "
        "demand, particularly for the Pro and Pro Max models. Average selling prices "
        "increased 4% to $988 per unit, offsetting a slight decline in unit shipments."
    )
    doc.add_paragraph(
        "Services revenue of $100.0 billion grew 12% year over year, driven by "
        "advertising, App Store commissions, and AppleCare+. The services gross margin "
        "expanded to 74.0% from 71.7%, reflecting operating leverage and a favorable "
        "revenue mix shift toward higher-margin services like advertising and licensing."
    )

    # Risk Factors
    doc.add_heading("Risk Factors", level=1)
    doc.add_paragraph(
        "Key risks include: (1) Regulatory headwinds in the EU under the Digital Markets "
        "Act, which could require App Store fee reductions estimated at $4-6 billion in "
        "annual revenue impact. (2) China market weakness, where revenue declined 8% in "
        "FY2024 due to competition from Huawei and macroeconomic softness. (3) AI "
        "execution risk — Apple Intelligence features must demonstrate clear consumer "
        "value to justify the company's late entry into generative AI."
    )

    # Valuation
    doc.add_heading("Valuation", level=1)
    doc.add_paragraph(
        "At current levels, AAPL trades at 30.2x trailing earnings and 28.5x forward "
        "earnings. While this represents a premium to the S&P 500, we believe it is "
        "justified by Apple's best-in-class capital return program ($110 billion "
        "returned to shareholders in FY2024), ecosystem stickiness (98% retention rate), "
        "and the optionality from Apple Intelligence."
    )

    # Capital Allocation
    doc.add_heading("Capital Allocation", level=1)
    doc.add_paragraph(
        "Apple returned $110.2 billion to shareholders in FY2024 through $94.9 billion "
        "in share repurchases and $15.2 billion in dividends. The company's stated goal "
        "of achieving net cash neutral suggests continued aggressive buybacks. Free cash "
        "flow of $108.3 billion provides ample capacity for capital returns while "
        "maintaining investments in R&D ($31.4 billion in FY2024)."
    )

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Apple remains a core holding for quality-oriented investors. The services "
        "transformation is well underway, and we expect this segment to contribute "
        "an increasing share of profits over the next 3-5 years. Near-term catalysts "
        "include the iPhone 17 cycle (expected Sep 2025) and the full rollout of "
        "Apple Intelligence features. We reiterate our OVERWEIGHT rating."
    )

    doc.save(path)
    print(f"  ✅ Created: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# 5. PDF — 10-K style document (MD&A and Risk Factors)
# ═══════════════════════════════════════════════════════════════════════════

def create_sample_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    path = os.path.join(DATA_DIR, "apple_10k_excerpt_2024.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Title"], fontSize=18, spaceAfter=20
    )
    heading_style = ParagraphStyle(
        "CustomHeading", parent=styles["Heading1"], fontSize=14, spaceAfter=12,
        spaceBefore=20
    )
    body_style = ParagraphStyle(
        "CustomBody", parent=styles["Normal"], fontSize=10, leading=14,
        spaceAfter=8
    )

    story = []

    # Cover page
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph("APPLE INC.", title_style))
    story.append(Paragraph("ANNUAL REPORT PURSUANT TO SECTION 13 OR 15(d)", body_style))
    story.append(Paragraph("OF THE SECURITIES EXCHANGE ACT OF 1934", body_style))
    story.append(Paragraph("For the fiscal year ended September 28, 2024", body_style))
    story.append(Paragraph("Commission File Number: 001-36743", body_style))
    story.append(PageBreak())

    # Item 7: MD&A
    story.append(Paragraph("ITEM 7. MANAGEMENT'S DISCUSSION AND ANALYSIS OF FINANCIAL CONDITION AND RESULTS OF OPERATIONS", heading_style))

    mda_paragraphs = [
        "The following discussion should be read in conjunction with the consolidated "
        "financial statements and notes thereto included in Part II, Item 8 of this "
        "Annual Report on Form 10-K.",

        "FISCAL 2024 HIGHLIGHTS: Total net revenue for fiscal 2024 was $394.3 billion, "
        "an increase of 2.9% compared to fiscal 2023. The increase was driven primarily "
        "by higher Services revenue and higher iPhone revenue, partially offset by lower "
        "iPad revenue. The Company's total gross margin percentage increased to 46.2% in "
        "fiscal 2024 from 44.1% in fiscal 2023, driven primarily by a favorable shift "
        "in revenue mix toward Services, which carry higher margins.",

        "PRODUCTS REVENUE: Products revenue was $294.3 billion in fiscal 2024, a decrease "
        "of 1% compared to fiscal 2023. iPhone revenue was $201.2 billion, representing "
        "51% of total net revenue. Mac revenue increased 2% to $29.9 billion, driven by "
        "the MacBook Pro with M3 chip. iPad revenue decreased 8% to $26.7 billion "
        "primarily due to the different timing of iPad launches year over year. Wearables, "
        "Home and Accessories revenue was $36.5 billion, down 3% year over year.",

        "SERVICES REVENUE: Services revenue was $100.0 billion in fiscal 2024, an increase "
        "of 12% compared to fiscal 2023. The growth was driven by increases in advertising, "
        "the App Store, and cloud services. The Company now has more than 1 billion paid "
        "subscriptions across its services offerings, an increase of 15% from the prior year.",

        "OPERATING EXPENSES: Total operating expenses were $57.5 billion in fiscal 2024, "
        "an increase of 5% compared to fiscal 2023. Research and development expenses "
        "increased 5% to $31.4 billion, reflecting continued investment in product "
        "innovation including Apple Intelligence and spatial computing. Selling, general "
        "and administrative expenses increased 5% to $26.1 billion.",

        "OPERATING INCOME: Operating income was $124.8 billion in fiscal 2024, an increase "
        "of 9% compared to fiscal 2023. Operating margin was 31.6% compared to 29.8% in "
        "the prior year, reflecting gross margin improvement from the revenue mix shift "
        "toward higher-margin Services.",

        "CASH FLOW: The Company generated operating cash flow of $118.3 billion in fiscal "
        "2024. Capital expenditures were $10.0 billion. Free cash flow was $108.3 billion. "
        "The Company returned $110.2 billion to shareholders through share repurchases of "
        "$94.9 billion and dividend payments of $15.2 billion.",
    ]

    for para in mda_paragraphs:
        story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 6))

    story.append(PageBreak())

    # Item 1A: Risk Factors
    story.append(Paragraph("ITEM 1A. RISK FACTORS", heading_style))

    risk_paragraphs = [
        "An investment in the Company involves risk. The following discussion of risk "
        "factors contains forward-looking statements. The Company's actual results may "
        "differ materially from those discussed in these forward-looking statements.",

        "GLOBAL AND MACROECONOMIC RISKS: The Company's operations and performance depend "
        "significantly on worldwide economic conditions. Uncertainty about global economic "
        "conditions, including inflation, interest rates, and geopolitical tensions, could "
        "cause consumers and businesses to postpone spending. The Company experienced "
        "revenue declines in Greater China of 8% in fiscal 2024, reflecting both "
        "macroeconomic softness and increased local competition.",

        "COMPETITION: The markets for the Company's products and services are highly "
        "competitive. The Company faces substantial competition from companies such as "
        "Samsung, Google, and Huawei in smartphones, and from Microsoft, Google, and "
        "Amazon in cloud and AI services. In China specifically, Huawei's resurgence with "
        "its Mate 60 series has intensified competitive pressures.",

        "SUPPLY CHAIN: The Company relies on single or limited sources for certain "
        "components, including certain semiconductors manufactured by Taiwan Semiconductor "
        "Manufacturing Company (TSMC). Disruptions in the supply chain, whether from "
        "natural disasters, geopolitical events, or other factors, could adversely affect "
        "the Company's ability to meet customer demand.",

        "REGULATORY AND LEGAL: The Company is subject to complex and evolving laws and "
        "regulations worldwide. The European Union's Digital Markets Act may require "
        "significant changes to the App Store's business model, potentially impacting "
        "annual App Store revenue by an estimated $4-6 billion. The Company is also "
        "subject to ongoing antitrust investigations in the United States and other "
        "jurisdictions.",

        "ARTIFICIAL INTELLIGENCE: The Company's ability to compete effectively depends "
        "increasingly on its AI capabilities. The Company launched Apple Intelligence "
        "in fiscal 2024, but faces significant competition from established AI leaders. "
        "Failure to deliver compelling AI features could negatively impact the Company's "
        "competitive position and ability to attract and retain customers.",

        "FOREIGN EXCHANGE: The Company's international operations expose it to fluctuations "
        "in foreign currency exchange rates. Approximately 60% of the Company's revenue is "
        "generated outside the United States. A stronger U.S. dollar negatively impacts "
        "reported international revenue. In fiscal 2024, foreign exchange movements had a "
        "negative impact of approximately $3.2 billion on total revenue.",
    ]

    for para in risk_paragraphs:
        story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    print(f"  ✅ Created: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating sample data files...\n")

    # Always available (no special deps)
    create_sample_txt()
    create_sample_csv()

    # Require pandas + openpyxl
    try:
        create_sample_xlsx()
    except ImportError as e:
        print(f"  ⚠️  Skipping XLSX (missing dependency: {e})")

    # Requires python-docx
    try:
        create_sample_docx()
    except ImportError as e:
        print(f"  ⚠️  Skipping DOCX (missing dependency: {e})")

    # Requires reportlab
    try:
        create_sample_pdf()
    except ImportError as e:
        print(f"  ⚠️  Skipping PDF (missing dependency: {e})")

    print("\nDone! Sample files are in the data/ directory.")
